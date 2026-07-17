from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_PATH = Path("Product_Report_Helmet_Detection.docx")


SECTIONS = [
    (
        "1. Thong Tin Chung",
        [
            "Ten de tai: Phat hien nguoi dieu khien xe co doi mu bao hiem hay khong bang mo hinh YOLO.",
            "Loai san pham: He thong nhan dien doi tuong tren anh giao thong.",
            "Muc tieu san pham: Ho tro nhan dien nguoi lai xe, phan biet doi mu bao hiem va khong doi mu bao hiem.",
        ],
    ),
    (
        "2. Mo Ta Bai Toan",
        [
            "Bai toan dat ra la tu mot anh giao thong, mo hinh can xac dinh va dinh vi cac doi tuong lien quan den bai toan an toan giao thong.",
            "Trong pham vi du an, nhom tap trung vao 3 class chinh: helmet, no_helmet va rider.",
            "Day la bai toan object detection vi he thong khong chi can phan loai ma con phai xac dinh vi tri doi tuong tren anh.",
        ],
    ),
    (
        "3. Du Lieu Su Dung",
        [
            "Bo du lieu local duoc luu trong thu muc data/ theo dinh dang YOLO.",
            "Tong cong bo du lieu co 942 anh, trong do 800 anh train va 142 anh valid.",
            "Ti le chia du lieu xap xi 85% train va 15% valid.",
            "Ban dau dataset co 5 class, nhung de phu hop muc tieu cuoi cung, nhom da remap va rut gon thanh 3 class.",
            "Bo du lieu hien tai khong co test split rieng, do do viec danh gia chu yeu dua tren tap valid.",
        ],
    ),
    (
        "4. Tien Xu Ly Va Khao Sat Du Lieu",
        [
            "Nhom tien hanh kiem tra cau truc train/images, train/labels, valid/images, valid/labels va file data.yaml.",
            "Qua trinh khao sat cho thay bo du lieu gom nhieu tinh huong giao thong thuc te, doi tuong nho, goc nhin kho va dieu kien anh khong dong deu.",
            "Ten class trong bo du lieu goc khong day du y nghia, vi vay nhom da phai phan tich va chuan hoa lai class phuc vu bai toan cuoi.",
            "Sau phan tich, nhom quyet dinh tap trung vao 3 class la helmet, no_helmet va rider de tang tinh phu hop cua mo hinh.",
        ],
    ),
    (
        "5. Mo Hinh Va Phuong Phap",
        [
            "Mo hinh chinh duoc su dung la YOLOv8.",
            "Ban dau, nhom xay dung baseline tren bo du lieu 5 class de co moc so sanh.",
            "Sau do, nhom fine-tune YOLOv8 tren bo du lieu 3 class da duoc remap.",
            "Nhieu thu nghiem bo sung cung duoc thuc hien nhu thay doi kich thuoc anh, oversample class no_helmet, fine-tune 2 stage, thu YOLO11 va huong detector + classifier.",
            "Ket qua cho thay YOLOv8 fine-tune tren bo 3 class van la lua chon hieu qua nhat trong pham vi du lieu hien co.",
        ],
    ),
    (
        "6. Qua Trinh Thu Nghiem",
        [
            "Thu nghiem 1: YOLOv8 baseline tren bai toan 5 class.",
            "Thu nghiem 2: Chuyen bai toan thanh 3 class va fine-tune lai mo hinh.",
            "Thu nghiem 3: Tinh chinh imgsz, epoch, patience va cac tham so augmentation.",
            "Thu nghiem 4: Thu nghiem YOLO11 de so sanh kien truc.",
            "Thu nghiem 5: Thu huong 2-stage detector + classifier, nhung huong nay chua vuot duoc mo hinh YOLOv8 toi uu do class head qua yeu.",
        ],
    ),
]


RESULT_PARAGRAPHS = [
    "Ket qua baseline YOLOv8 ban dau tren bai toan 5 class dat mAP50 = 65.1% va mAP50-95 = 30.2%.",
    "Sau khi remap va fine-tune YOLOv8 tren bai toan 3 class, mo hinh tot nhat dat standard validation voi Precision = 71.0%, Recall = 78.4%, mAP50 = 77.4% va mAP50-95 = 36.8%.",
    "Khi su dung TTA trong giai doan danh gia, mo hinh dat mAP50 = 77.9% va mAP50-95 = 36.8%.",
    "Theo tung class o ket qua standard validation, helmet dat mAP50 = 76.3%, no_helmet dat mAP50 = 65.6% va rider dat mAP50 = 90.4%.",
    "Class rider la class de nhan biet nhat, trong khi no_helmet la class kho nhat do doi tuong nho, de bi che khuat va co nhieu truong hop goc nhin phuc tap.",
]


COMPARISON_ROWS = [
    ("YOLOv8 baseline ban dau", "5 class", "65.1%", "30.2%"),
    ("YOLOv8 fine-tune", "3 class", "77.4%", "36.8%"),
    ("YOLOv8 fine-tune + TTA", "3 class", "77.9%", "36.8%"),
]


FINAL_SECTIONS = [
    (
        "8. Nhan Xet Va Danh Gia",
        [
            "Qua trinh fine-tune da giup mo hinh tang khoang 12.3 diem mAP50 so voi baseline ban dau.",
            "Huong 3 class cho ket qua tot hon 5 class vi phu hop hon voi muc tieu cuoi cung cua bai toan.",
            "YOLO11 va huong 2-stage duoc thu nghiem de mo rong bai toan, tuy nhien trong bo du lieu hien tai van chua vuot duoc YOLOv8 toi uu.",
            "Dieu nay cho thay voi bai toan nay, chat luong va cau truc du lieu anh huong rat lon den gioi han cua mo hinh.",
        ],
    ),
    (
        "9. Uu Diem Cua San Pham",
        [
            "Mo hinh hoat dong duoc tren anh giao thong thuc te va co the xuat bbox kem nhan du doan.",
            "Do chinh xac tren class rider rat tot, giup mo hinh nhan dien nguoi lai xe on dinh.",
            "Pipeline YOLOv8 de huan luyen, de trien khai va phu hop cho demo du an.",
        ],
    ),
    (
        "10. Han Che Cua San Pham",
        [
            "Bo du lieu khong co test split rieng, nen danh gia chua hoan toan khach quan.",
            "Class no_helmet van kho hoc va co do chinh xac thap hon cac class con lai.",
            "Du lieu co nhieu anh vat the nho, anh mo, watermark va sai khac mien du lieu, lam giam kha nang tong quat hoa.",
        ],
    ),
    (
        "11. Huong Phat Trien",
        [
            "Bo sung du lieu moi va lam sach nhan, dac biet cho class no_helmet.",
            "Tao bo test rieng de danh gia khach quan hon.",
            "Thu nghiem them huong crop theo vung dau hoac classifier bo tro cho truong hop kho.",
            "Toi uu hoa suy luan de trien khai tren video hoac he thong camera thuc te.",
        ],
    ),
    (
        "12. Ket Luan",
        [
            "Du an da xay dung thanh cong mo hinh nhan dien doi mu bao hiem dua tren YOLO.",
            "Mo hinh YOLOv8 fine-tune tren bo du lieu 3 class la ket qua tot nhat cua nhom trong qua trinh thu nghiem.",
            "San pham dat mAP50 gan 78%, the hien tiem nang ung dung thuc te trong bai toan ho tro giam sat an toan giao thong.",
            "Neu co them du lieu va thoi gian toi uu, san pham co the tiep tuc duoc nang cap de dat do chinh xac cao hon.",
        ],
    ),
]


def add_heading(document: Document, text: str, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BAO CAO SAN PHAM DU AN NHAN DIEN DOI MU BAO HIEM")
    title_run.bold = True
    title_run.font.size = Pt(18)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Mo hinh YOLO cho bai toan giao thong").italic = True

    document.add_paragraph("")

    intro = document.add_paragraph()
    intro.add_run("Tom tat: ").bold = True
    intro.add_run(
        "Bao cao nay trinh bay qua trinh xay dung, thu nghiem va danh gia mo hinh nhan dien "
        "nguoi dieu khien xe co doi mu bao hiem hay khong dua tren kien truc YOLO."
    )

    for heading, bullets in SECTIONS:
        add_heading(document, heading)
        add_bullets(document, bullets)

    add_heading(document, "7. Ket Qua Dat Duoc")
    for paragraph_text in RESULT_PARAGRAPHS:
        document.add_paragraph(paragraph_text)

    table_title = document.add_paragraph()
    table_title.add_run("Bang so sanh ket qua giua baseline va mo hinh fine-tune:").bold = True

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Mo hinh"
    header_cells[1].text = "Bai toan"
    header_cells[2].text = "mAP50"
    header_cells[3].text = "mAP50-95"

    for row in COMPARISON_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    for heading, bullets in FINAL_SECTIONS:
        add_heading(document, heading)
        add_bullets(document, bullets)

    final_note = document.add_paragraph()
    final_note.add_run("Goi y bo sung khi nop bao cao: ").bold = True
    final_note.add_run(
        "Ban nen chen them anh dataset mau, anh ket qua predict cua mo hinh, "
        "bang confusion matrix hoac PR curve neu can minh hoa them phan ket qua."
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH.resolve()}")


if __name__ == "__main__":
    main()
