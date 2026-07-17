from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("Project_Presentation_Slides_Content.docx")


SLIDES = [
    {
        "title": "Slide 1. Tieu De De Tai",
        "points": [
            "De tai: Phat hien nguoi dieu khien xe co doi mu bao hiem hay khong bang Computer Vision.",
            "Huong thuc hien: Xay dung mo hinh nhan dien doi tuong tren anh giao thong bang YOLO.",
            "Muc tieu chinh: Phan biet 3 nhom doi tuong gom helmet, no_helmet, rider.",
        ],
        "notes": [
            "O slide mo dau, ban gioi thieu ngan gon ten de tai va bai toan thuc te ma nhom giai quyet.",
            "Nhan manh rang bai toan nay co y nghia trong an toan giao thong va giam sat tu dong.",
        ],
    },
    {
        "title": "Slide 2. Bai Toan Va Muc Tieu",
        "points": [
            "Bai toan dat ra: Tu anh giao thong, xac dinh nguoi lai xe co doi mu bao hiem hay khong.",
            "Input: Anh hoac khung hinh giao thong.",
            "Output mong muon: Vi tri doi tuong va nhan helmet, no_helmet, rider.",
            "Tieu chi danh gia: Precision, Recall, mAP50, mAP50-95.",
        ],
        "notes": [
            "Neu can noi don gian, ban co the noi day la bai toan object detection trong boi canh giao thong.",
            "Muc tieu cuoi cung khong phai detect tat ca vat the, ma la phuc vu bai toan doi mu bao hiem.",
        ],
    },
    {
        "title": "Slide 3. Du Lieu Su Dung",
        "points": [
            "Bo du lieu goc theo dinh dang YOLO, luu trong thu muc data/.",
            "Tong cong 942 anh: 800 anh train va 142 anh valid.",
            "Du lieu ban dau co 5 class, sau do duoc rut gon lai 3 class phuc vu bai toan cuoi.",
            "Cau truc du lieu: train/images, train/labels, valid/images, valid/labels.",
        ],
        "notes": [
            "Ban co the noi them rang bo du lieu khong co test split rieng, nen viec danh gia chu yeu dua tren valid.",
            "Neu thuyet trinh, nen nhac day la mot han che cua bai toan.",
        ],
    },
    {
        "title": "Slide 4. Tien Xu Ly Va Khao Sat Du Lieu",
        "points": [
            "Kiem tra dinh dang nhan YOLO va so luong anh, label cua tung split.",
            "Phan tich nhanh cho thay bo du lieu co nhieu anh giao thong thuc te, kich thuoc vat the nho va goc nhin da dang.",
            "Nhan thay bo 5 class khong toi uu cho muc tieu cuoi, vi vay nhom chuyen sang huong 3 class.",
            "Ba class cuoi cung duoc chon la: helmet, no_helmet, rider.",
        ],
        "notes": [
            "Day la slide de ban cho thay nhom khong train ngay ma co buoc hieu du lieu truoc.",
            "Neu co anh minh hoa, nen chen 2 den 4 anh mau co bbox.",
        ],
    },
    {
        "title": "Slide 5. Huong Giai Quyet",
        "points": [
            "Mo hinh nen tang duoc chon la YOLOv8 de thuc hien object detection.",
            "Ban dau thu nghiem tren bo 5 class de co baseline.",
            "Sau do toi uu lai thanh bai toan 3 class de tap trung vao muc tieu doi mu bao hiem.",
            "Ngoai ra, nhom cung thu nghiem them huong 2-stage va YOLO11 de so sanh.",
        ],
        "notes": [
            "Ban nen nhan manh ly do chon YOLOv8: phu hop object detection thoi gian gan thuc, de fine-tune va co tai lieu ho tro tot.",
            "Noi ro rang nhom da co qua trinh thu nghiem, so sanh, khong chi dung mot mo hinh duy nhat.",
        ],
    },
    {
        "title": "Slide 6. Cac Thu Nghiem Da Thuc Hien",
        "points": [
            "Thu nghiem 1: Train YOLO tren bo 5 class de co moc ket qua ban dau.",
            "Thu nghiem 2: Remap du lieu ve 3 class va fine-tune YOLOv8.",
            "Thu nghiem 3: Dieu chinh kich thuoc anh, oversample no_helmet, fine-tune 2 stage.",
            "Thu nghiem 4: Thu YOLO11 va huong 2-stage detector + classifier, nhung ket qua khong vuot qua YOLOv8 toi uu.",
        ],
        "notes": [
            "Slide nay giup thay rang nhom da co tu duy toi uu mo hinh va so sanh nghiem tuc.",
            "Neu can ngan gon, ban co the noi rang nhom uu tien huong nao cho ket qua thuc te tot nhat.",
        ],
    },
    {
        "title": "Slide 7. Ket Qua Mo Hinh Tot Nhat",
        "points": [
            "Mo hinh tot nhat la YOLOv8 3 class sau qua trinh fine-tune va tinh chinh tham so.",
            "Ket qua standard validation: mAP50 = 77.4%, mAP50-95 = 36.8%.",
            "Ket qua khi dung TTA luc danh gia: mAP50 = 77.9%, mAP50-95 = 36.8%.",
            "Theo tung class o standard val: helmet = 76.3%, no_helmet = 65.6%, rider = 90.4% theo mAP50.",
        ],
        "notes": [
            "Day la slide quan trong nhat, ban nen doc cham va ro cac chi so.",
            "Neu can, ban co the giai thich TTA la mot cach tang do on dinh khi danh gia bang cach du doan tren nhieu bien the cua anh.",
        ],
    },
    {
        "title": "Slide 8. Phan Tich Ket Qua",
        "points": [
            "Class rider dat ket qua cao nhat vi kich thuoc lon hon va de nhan biet hon.",
            "Class no_helmet la class kho nhat, do vat the nho, goc nhin kho, anh mo va nhieu truong hop bi che khuat.",
            "Mo hinh 3 class tot hon mo hinh 5 class vi phu hop hon voi muc tieu bai toan.",
            "YOLO11 va pipeline 2-stage duoc thu nghiem nhung chua vuot duoc YOLOv8 toi uu trong bo du lieu hien tai.",
        ],
        "notes": [
            "Slide nay nen dung de ly giai vi sao mo hinh dat nhu vay, tranh chi dua so lieu ma khong phan tich.",
            "Neu giang vien hoi tai sao chon 3 class, day la slide de tra loi.",
        ],
    },
    {
        "title": "Slide 9. Uu Diem Va Han Che",
        "points": [
            "Uu diem: Pipeline YOLOv8 hoat dong on dinh, de suy luan, co ket qua kha tot tren bai toan 3 class.",
            "Uu diem: Mo hinh dat mAP50 gan 78%, rider dat tren 90% theo mAP50.",
            "Han che: Bo du lieu khong co test split rieng, ten class ban dau khong day du, no_helmet van la class kho.",
            "Han che: Nhieu anh co vat the nho, anh mo, goc nhin xau va do lech mien du lieu.",
        ],
        "notes": [
            "Ban nen noi thang thac han che de bai thuyet trinh trong thuc va co chieu sau hon.",
            "Neu duoc hoi ve do chinh xac 90%, ban co the noi ket qua hien tai da duoc toi uu trong dieu kien du lieu san co, nhung van bi gioi han boi chat luong va cau truc du lieu.",
        ],
    },
    {
        "title": "Slide 10. Huong Phat Trien",
        "points": [
            "Bo sung va lam sach them du lieu cho class no_helmet.",
            "Tao test set rieng de danh gia khach quan hon.",
            "Thu nghiem them pipeline head-focused hoac classifier bo tro cho no_helmet.",
            "Toi uu mo hinh de trien khai tren video hoac camera thuc te theo thoi gian thuc.",
        ],
        "notes": [
            "Slide nay cho thay du an van co kha nang phat trien tiep.",
            "Ban nen chot rang huong cai thien quan trong nhat la du lieu, dac biet voi class no_helmet.",
        ],
    },
    {
        "title": "Slide 11. Ket Luan",
        "points": [
            "Nhom da xay dung duoc he thong nhan dien doi mu bao hiem dua tren YOLO.",
            "Huong 3 class cho ket qua tot hon huong 5 class va cac thu nghiem thay the khac.",
            "Mo hinh YOLOv8 fine-tune la lua chon phu hop nhat de demo va bao cao du an hien tai.",
            "Du an co tinh ung dung thuc te va co the tiep tuc cai tien neu co them du lieu va thoi gian.",
        ],
        "notes": [
            "O slide cuoi, ban nen tong ket gon va nhan manh gia tri thuc te cua de tai.",
            "Co the ket thuc bang cau: Em xin cam on thay co va cac ban da lang nghe.",
        ],
    },
]


def add_bullets(document: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        document.add_paragraph(item, style=style)


def build_document() -> Document:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Noi Dung Slide Thuyet Trinh Du An")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("De tai: Phat hien doi mu bao hiem bang YOLO").italic = True

    document.add_paragraph("")

    intro = document.add_paragraph()
    intro.add_run("Huong dan su dung: ").bold = True
    intro.add_run(
        "Moi muc ben duoi tuong ung voi 1 slide. "
        "Ban co the copy nguyen van phan 'Noi dung tren slide' sang PowerPoint, "
        "sau do doc theo phan 'Goi y loi trinh bay'."
    )

    for slide in SLIDES:
        document.add_page_break()

        heading = document.add_paragraph()
        heading_run = heading.add_run(slide["title"])
        heading_run.bold = True
        heading_run.font.size = Pt(16)

        body_label = document.add_paragraph()
        body_label.add_run("Noi dung tren slide:").bold = True
        add_bullets(document, slide["points"])

        notes_label = document.add_paragraph()
        notes_label.add_run("Goi y loi trinh bay:").bold = True
        add_bullets(document, slide["notes"])

    closing = document.add_page_break()
    _ = closing
    final_note = document.add_paragraph()
    final_note.add_run("Goi y bo sung hinh anh cho slide: ").bold = True
    final_note.add_run(
        "Nen chen them 1 anh dataset mau, 1 anh ket qua bbox, "
        "1 bang so sanh giua 5-class va 3-class, va 1 slide demo ket qua tot nhat."
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH.resolve()}")


if __name__ == "__main__":
    main()
